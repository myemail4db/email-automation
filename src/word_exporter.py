from pathlib import Path

from docx import Document
from docx.shared import Pt

from src.text_exporter import build_safe_stem
from src.text_filter import clean_email_body
from src.utils.date_utils import format_email_date


def add_separator(doc, char="=", length=60):
    doc.add_paragraph(char * length)



def add_header_field(doc, label, value):
    p = doc.add_paragraph()
    run_label = p.add_run(f"{label:<18}: ")
    run_label.bold = True
    p.add_run(str(value))



def save_email_to_word(subject, sender, date, body, output_dir, status="processed_review", format_type="word"):
    output_dir = Path(output_dir)
    output_dir.mkdir(parents=True, exist_ok=True)

    safe_stem = build_safe_stem(subject or "")
    file_path = output_dir / f"{safe_stem}.docx"

    counter = 1
    while file_path.exists():
        file_path = output_dir / f"{safe_stem}_{counter}.docx"
        counter += 1

    cleaned_body = clean_email_body(body)

    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    add_separator(doc)

    title = doc.add_paragraph()
    title_run = title.add_run("EMAIL DETAILS")
    title_run.bold = True
    title_run.font.size = Pt(12)

    add_separator(doc)

    add_header_field(doc, "Subject", subject or "No Subject")
    add_header_field(doc, "From", sender or "Unknown Sender")
    add_header_field(doc, "Date", format_email_date(date))
    add_header_field(doc, "Workflow Status", status)
    add_header_field(doc, "Export Format", format_type)
    add_header_field(doc, "File Name", file_path.name)

    add_separator(doc)
    doc.add_paragraph()

    for line in cleaned_body.split("\n"):
        doc.add_paragraph(line)

    doc.save(file_path)
    return str(file_path)
